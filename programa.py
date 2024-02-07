import csv
import os
import re
import time
import tkinter as tk
import logging
from tkinter import ttk
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
from pathlib import Path
from tkinter import messagebox
from tkinter import PhotoImage

###################################################################logging###########################################################
# Configuración básica del logger
num = 0
aux = 0
tabla = None
estado = True

directorio_actual = "C:\\Reportes\\Logs"
if not os.path.exists(directorio_actual):
    os.makedirs(directorio_actual)
archivo_csv = os.path.join(directorio_actual, 'registro.csv')
archivo_log = os.path.join(directorio_actual, 'registro.log')
if not os.path.isfile(archivo_log):
    with open(archivo_log, 'w'):
        pass  # Crea el archivo si no existe
logging.basicConfig(filename='registro.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
# Crear el archivo CSV y escribir el encabezado si el archivo no existe
if not os.path.isfile(archivo_csv):
    with open(archivo_csv, 'w', newline='') as csvfile:
        csvwriter = csv.writer(csvfile)
        csvwriter.writerow(['Serial', 'Modelo', 'Firmware', 'CPU', 'Memoria', 'Estado Firmware', 'Estado Tx', 'Estado Rx','Estado ONT'])
def borrar_productos():
    global tree
    for item in tree.get_children():
        tree.delete(item)

def ejecutar_Huawei():
    try:
        global num
        global tabla
        if num != 0:
            for item in tabla.get_children():
                tabla.delete(item)
           
            num = 0
            
        #Configuración del navegador y URL
        
        '''driver = webdriver.Chrome()
        driver.get("http://192.168.18.1")
        # Crear la carpeta para los Reportes
        ruta_carpeta1 = "C:\\Reportes"
        if not os.path.exists(ruta_carpeta1):
            os.makedirs(ruta_carpeta1)
        #Iniciar Sesión
        username = WebDriverWait(driver,7).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "input[name='txt_Username']")))
        password = WebDriverWait(driver,7).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "input[name='txt_Password']")))
        username.clear()
        username.send_keys("Epadmin")
        password.clear()
        password.send_keys("adminEp")
        #Botón iniciar sesión
        button =WebDriverWait(driver,7).until(EC.element_to_be_clickable((By.ID,"loginbutton"))).click()
        #Botón Para ir al Device Information
        button =WebDriverWait(driver,7).until(EC.element_to_be_clickable((By.ID,"icon_Systeminfo"))).click()
        time.sleep(5)
        # Esperar a que el iframe esté presente antes de cambiar al contexto del iframe
        iframe_element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, "menuIframe")))
        # Cambiar al contexto del iframe
        driver.switch_to.frame(iframe_element)
        # Encuentra el elemento que contiene el texto que quieres extraer
        serial_ont = driver.find_element(By.XPATH,'//*[@id="td3_2"]')
        modelo_ont = driver.find_element(By.XPATH,'//*[@id="td1_2"]')
        firmware_ont = driver.find_element(By.XPATH,'//*[@id="td5_2"]')
        cpu_ont = driver.find_element(By.XPATH,'//*[@id="td9_2"]')
        mem_ont = driver.find_element(By.XPATH,'//*[@id="td10_2"]')
        serial_extraido = str(serial_ont.text)
        modelo_extraido = str(modelo_ont.text)
        firmware_extraido = str(firmware_ont.text)
        cpu_extraido = str(cpu_ont.text)
        mem_extraido = str(mem_ont.text)
        # Crea una carpeta con el nombre extraído
        carpeta_nombre = f"C:\\Reportes\\{serial_extraido}"
        os.makedirs(carpeta_nombre, exist_ok=True)
        # Guarda el screenshot en la carpeta creada
        ruta_screenshot_deviceinfo = os.path.join(carpeta_nombre, "Device_Info.png")
        driver.get_screenshot_as_file(ruta_screenshot_deviceinfo)
        # Volver al contexto principal (fuera del iframe)
        driver.switch_to.default_content()
        #Botón Para ir al Optical info
        button =WebDriverWait(driver,7).until(EC.element_to_be_clickable((By.ID,"name_opticinfo"))).click()
        # Esperar a que el iframe esté presente antes de cambiar al contexto del iframe
        iframe_element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, "menuIframe")))
        time.sleep(2)
        # Cambiar al contexto del iframe
        driver.switch_to.frame(iframe_element)
        # Encuentra el elemento que contiene el texto que quieres extraer
        tx_ont = driver.find_element(By.XPATH,'//*[@id="optic_status_table"]/tbody/tr[6]/td[2]')
        rx_ont = driver.find_element(By.XPATH,'//*[@id="optic_status_table"]/tbody/tr[7]/td[2]')
        rx_extraido = rx_ont.text
        tx_extraido = tx_ont.text
        tx_limpio = float(re.search(r'\d+', tx_extraido).group()) if re.search(r'\d+', tx_extraido) else 0
        rx_limpio = float(re.search(r'\d+', rx_extraido).group()) if re.search(r'\d+', rx_extraido) else -1
        #Tomar Screenshot de Optical Info
        ruta_screenshot_opticalinfo = os.path.join(carpeta_nombre, "Optical_Info.png")
        driver.get_screenshot_as_file(ruta_screenshot_opticalinfo)
        # Volver al contexto principal (fuera del iframe)
        driver.switch_to.default_content()
        #Botón Para ir al Ethernet Status
        button =WebDriverWait(driver,7).until(EC.element_to_be_clickable((By.ID,"name_ethinfo"))).click()
        time.sleep(2)
        #Tomar Screenshot de Ethernet Status
        ruta_screenshot_ethinfo = os.path.join(carpeta_nombre, "Eth_info.png")
        driver.get_screenshot_as_file(ruta_screenshot_ethinfo)
        #Botón Para ir al Ethernet Status
        button =WebDriverWait(driver,7).until(EC.element_to_be_clickable((By.ID,"name_wlaninfo"))).click()
        time.sleep(2)
        # Esperar a que el iframe esté presente antes de cambiar al contexto del iframe
        iframe_element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, "menuIframe")))
        # Cambiar al contexto del iframe
        driver.switch_to.frame(iframe_element)
        time.sleep(2)
        button2 =WebDriverWait(driver,7).until(EC.element_to_be_clickable((By.XPATH, '//*[@id="btn_health_check"]'))).click()
        #Tomar Screenshot de Wireless
        time.sleep(7)
        ruta_screenshot_wlaninfo = os.path.join(carpeta_nombre, "Wlan_info.png")
        driver.get_screenshot_as_file(ruta_screenshot_wlaninfo)
        ###################################################Documento##################################################################################
        datos = {
            'SERIAL': serial_extraido,
            'MODELO': modelo_extraido,
            'FIRMWARE': firmware_extraido,
            'CPU': cpu_extraido,
            'MEM': mem_extraido
        }
        # Cargar el documento existente
        documento = Document('C:\\Reportes\\Plantilla_Reporte.docx')  # Reemplaza con la ruta de tu documento
        # Limpiar el texto extraído para asegurarse de que sea un nombre de carpeta y archivo válido
        texto_extraido_limpio = ''.join(caracter for caracter in serial_extraido if caracter.isalnum() or caracter in ['_', ' ', '-', '(' , ')'])
        # Construir la ruta completa para la imagen
        imagen_ruta = f'C:\\Reportes\\{texto_extraido_limpio}\\Device_Info.png'
        imagen_ruta2 = f'C:\\Reportes\\{texto_extraido_limpio}\\Optical_Info.png'
        imagen_ruta3 = f'C:\\Reportes\\{texto_extraido_limpio}\\Eth_info.png'
        imagen_ruta4 = f'C:\\Reportes\\{texto_extraido_limpio}\\Wlan_info.png'
        # Añadir la imagen al documento con tamaño y posición personalizados
        ancho_imagen = Inches(4.0)
        altura_imagen = Inches(4.0)
        # Agregar una imagen al documento
        for key, value in datos.items():
            for paragraph in documento.paragraphs:
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))
        documento.add_picture(imagen_ruta, width=ancho_imagen, height=altura_imagen)
        documento.add_picture(imagen_ruta2, width=ancho_imagen, height=altura_imagen)
        documento.add_picture(imagen_ruta3, width=ancho_imagen, height=altura_imagen)
        documento.add_picture(imagen_ruta4, width=ancho_imagen, height=altura_imagen)
        # Guardar el documento modificado
        ruta_guardado = f'C:\\Reportes\\{texto_extraido_limpio}\\{serial_extraido}.docx'
        documento.save(ruta_guardado)
        #########################################Validadores###########################################################
        txmin = 0.5
        txmax = 5
        rxmin = -27
        rxmax = -8
        estado_tx = ""
        estado_rx = ""
        estado_firmware = ""
        dispositivostatus =""
        if  txmin <= tx_limpio <= txmax:
            estado_tx = "Potencia óptica TX Correcta"
        else:
            estado_tx = "Potencia óptica TX Incorrecta"
        if  rxmin <= rx_limpio <= rxmax:
            estado_rx = "Potencia óptica RX Correcta"
            etiqueta = tk.Label(root, text=f"<font color='red'>{estado_rx}</font>", font=("Arial", 12))
        else:
            estado_rx = "Potencia óptica RX Incorrecta"
        if  firmware_extraido ==  "V5R020C10S115":
            estado_firmware = "El Firmware está actualizado"
        else:
            estado_firmware = "El Firmware no está actualizado"
        if estado_firmware == "El Firmware está actualizado" and estado_tx == "Potencia óptica TX Correcta" and estado_rx == "Potencia óptica RX Correcta":
            dispositivostatus = "☺ -  Dispositivo aprobado"
            etiqueta = tk.Label(root, text=dispositivostatus, fg="green", font=("Arial", 16, "bold"))
        else:
            dispositivostatus = "☹ - Revisar dispositivo"
            etiqueta = tk.Label(root, text=dispositivostatus, fg="red", font=("Arial", 16, "bold"))
        etiqueta = tk.Label(root, text=dispositivostatus)
        if dispositivostatus == "☺ -  Dispositivo aprobado":
            etiqueta = tk.Label(root, text=dispositivostatus, fg="green", font=("Arial", 20, "bold"))
        else:
            etiqueta = tk.Label(root, text=dispositivostatus, fg="red", font=("Arial", 20, "bold"))
        etiqueta.pack()
        '''
        ########################################Tabla##################################################################
        frame_tabla = tk.Frame(root)
        frame_tabla.pack()
        # Create a table and insert data
        global estado
        
        if estado == True :
            tabla = ttk.Treeview(root)
        
        global aux
        
        tabla["columns"] = ("Atributo", "Valor")
        tabla.column("#0", stretch=tk.NO, width=0)  # Oculta la primera columna
        
        # Estilo para personalizar la tabla
        estilo = ttk.Style()
        estilo.configure("Treeview",
                 background="#FFFFFF",  # Color de fondo de la tabla
                 foreground="black",   # Color del texto
                 rowheight=25)          # Altura de la fila
        estilo.configure("Treeview.Heading", font=("Arial", 10, "bold"))
        tabla.tag_configure("oddrow", background="white")  # Color de fondo de filas impares
        tabla.tag_configure("evenrow", background="#F0F0F0")  # Color de fondo de filas pares
        for columna in tabla["columns"]:
                tabla.column("Atributo", anchor=tk.W, width=120)
                tabla.column("Valor", anchor=tk.W, width=200)
                tabla.heading("Atributo", text="Atributo", anchor=tk.W)
                tabla.heading("Valor", text="Valor", anchor=tk.W)
        tabla.insert("", tk.END, values=("Estado TX",aux))
        tabla.insert("", tk.END, values=("Estado RX", '''estado_rx'''))
        tabla.insert("", tk.END, values=("Estado Firmware", '''estado_firmware'''))
        tabla.insert("", tk.END, values=("Serial", '''serial_extraido'''))
        tabla.insert("", tk.END, values=("Modelo", '''modelo_extraido'''))
        tabla.insert("", tk.END, values=("Firmware", '''firmware_extraido'''))
        tabla.insert("", tk.END, values=("CPU", '''cpu_extraido'''))
        tabla.insert("", tk.END, values=("Memoria", '''mem_extraido'''))
        tabla.insert("", tk.END, values=("Potencia TX", '''tx_extraido'''))
        tabla.insert("", tk.END, values=("Potencia RX", '''rx_extraido'''))
       
        if(estado == True):
            tabla.pack(pady=0)
            estado = False
        print(tabla)
        
        
        num = num + 1
        aux = aux + 1

        
    except Exception as e:
        messagebox.showerror("Error", f"Se produjo un error: {str(e)}")
        #Datos para Log
        
        #logging.info(f'Serial: ' '''{serial_extraido}''')
        #logging.info(f'Modelo: ' '''{modelo_extraido}''')
        #logging.info(f'Firmware:'  '''{firmware_extraido}''')
        #logging.info(f'CPU:'  '''{cpu_extraido}''')
        #logging.info(f'Memoria:'  '''{mem_extraido}''')
        #logging.info(f'FirmwareStatus:'  '''{estado_firmware}''')
        #logging.info(f'RXOptical:' ''' {estado_rx}''')
        #logging.info(f'TXOptical:'  '''{estado_tx}''')
        #logging.info(f'StatusONT:'  '''{dispositivostatus}''')
    #with open('registro.csv', 'w', encoding='utf-8') as csvfile:
    #    csvwriter = csv.writer(csvfile)
    #    try:
    #        csvwriter.writerow([serial_extraido, modelo_extraido, firmware_extraido, cpu_extraido, mem_extraido,estado_firmware,estado_tx,estado_rx,dispositivostatus])
    #    except UnicodeEncodeError:
    #        pass
    ## Crear la ventana principal
def salir():
    respuesta = messagebox.askyesno("Salir", "¿Estás seguro que quieres salir?")
    if respuesta:
        
        root.destroy()  # Cerrar la ventana principal
root = tk.Tk()
root.title("Generador de Diagnosticos de ONT - Logistica Inversa")
root.geometry("500x800")
root.iconbitmap("logo.ico")
# Cargar la imagen de fondo
imagen_fondo = PhotoImage(file="fondo.png")
# Crear un Label con la imagen de fondo
label_fondo = tk.Label(root, image=imagen_fondo)
label_fondo.place(relwidth=1, relheight=1)
# Botón para ejecutar el Diagnostico Huawei
btn_ejecutar = tk.Button(root, text="Generar informe", command=ejecutar_Huawei, font=("Arial", 10, "bold"))

btn_ejecutar.pack(pady=0, side="top")
# Botón para salir
btn_salir = tk.Button(root, text="Salir del programa", command=salir, font=("Arial", 10, "bold"))
btn_salir.pack(pady=10, side="top")
# Iniciar el bucle principal de la interfaz gráfica
root.mainloop()

